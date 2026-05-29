#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
官方文档格式化脚本 - v7 中英文混排优化版
严格遵循用户提供的正确格式样例
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def fix_quotes(text):
    """将英文双引号替换为中文弯引号"""
    # 规则：遇到" 如果是开引号（前面是空格或在行首），替换为“，否则替换为”
    result = ''
    i = 0
    open_quote = True
    while i < len(text):
        if text[i] == '"':
            if open_quote:
                result += '\u201c'  # 中文左引号
                open_quote = False
            else:
                result += '\u201d'  # 中文右引号
                open_quote = True
            i += 1
        elif text[i] == "'":
            # 单引号也一并处理
            if open_quote:
                result += '\u2018'
                open_quote = False
            else:
                result += '\u2019'
                open_quote = True
            i += 1
        else:
            result += text[i]
            # 如果前面不是字符，下一个引号算作开引号
            if text[i].isspace() or i == len(text) - 1:
                open_quote = True
            i += 1
    return result

def split_text(text):
    """将文本拆分为中文和非中文（数字/英文）部分"""
    # 先修正引号
    text = fix_quotes(text)
    parts = []
    current = ''
    is_ascii = None
    
    for char in text:
        char_is_ascii = bool(re.match(r'[0-9a-zA-Z\.\-\:\/\,\%\s\+\(\)]+', char))
        if is_ascii is None:
            is_ascii = char_is_ascii
            current = char
        elif is_ascii == char_is_ascii:
            current += char
        else:
            parts.append((current, is_ascii))
            current = char
            is_ascii = char_is_ascii
    
    if current:
        parts.append((current, is_ascii))
    return parts

def format_official_document(input_path, output_path, doc_type='general'):
    """格式化官方文档"""
    doc = Document(input_path)
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    
    # 格式化段落（保留原有段落结构和空行，只格式化样式）
    for para in doc.paragraphs:
        text = para.text
        # 保留空行，不删除
        if not text.strip():
            # 空行也设置行间距
            para.paragraph_format.line_spacing = Pt(28)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue
        
        # 重置样式
        para.style = doc.styles['Normal']
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(28)
        
        # 清空段落内容，重新添加（处理中英文字体）
        # 先保存文字，然后重建
        original_text = para.text
        for run in para.runs:
            run.text = ''
        
        # ===== 检测各种标题类型 =====
        is_title = False
        chinese_font = '仿宋_GB2312'
        text_stripped = text.strip()
        
        # 1. 附件标识（附件1）
        if doc_type == '附件' and re.match(r'^附件\d+$', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            chinese_font = '黑体'
            is_title = True
            para.add_run(original_text).bold = True
        
        # 2. 大标题
        elif (any(k in text_stripped for k in ['请示', '报告', '通知', '方案', '总结', '规定', '办法', '意见', '函', '纪要', '明细表']) 
              and len(text_stripped) < 50 and len(text_stripped) > 10 
              and '，' not in text_stripped and '。' not in text_stripped
              and not re.match(r'^[一二三四五六七八九十]+、', text_stripped) 
              and not re.match(r'^（[一二三四五六七八九十]+）', text_stripped) 
              and not re.match(r'^\d+\.', text_stripped)):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            chinese_font = '方正小标宋简体'
            is_title = True
            # 大标题使用二号字体（22pt）
            run = para.add_run(original_text)
            run.font.name = '方正小标宋简体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
            run.font.size = Pt(22)
            run.bold = False
            # 跳过后续的字体设置
            continue
        
        # 3. 称谓
        elif text_stripped.endswith('：') and any(k in text_stripped for k in ['领导', '公司', '局', '委', '办', '同志']):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 4. 一级标题：一、XXX
        elif re.match(r'^[一二三四五六七八九十]+、', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '黑体'
            is_title = True
            run = para.add_run(original_text)
            run.bold = False  # 明确设置为不加粗
        
        # 5. 二级标题：（一）XXX
        elif re.match(r'^（[一二三四五六七八九十]+）', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '楷体_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 6. 三级标题：1.XXX
        elif re.match(r'^\d+\.\s*\S', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 7. 四级标题：（1）XXX
        elif re.match(r'^（\d+）', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 8. 五级标题：① XXX
        elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text_stripped):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 9. 附件说明
        elif text_stripped.startswith('附件：') or text_stripped.startswith('附：'):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0.74)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 10. 附件列表项（2.XXX、3.XXX）
        elif (re.match(r'^[2-9]\.', text_stripped) or re.match(r'^\d+\.\s*\S', text_stripped)) and len(text_stripped) < 30:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0.74)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 11. 结束语
        elif '妥否' in text_stripped and '请批示' in text_stripped:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(0.74)
            chinese_font = '仿宋_GB2312'
            is_title = True
            para.add_run(original_text)
        
        # 12. 落款和日期
        elif doc_type == '请示' or doc_type == '函':
            is_signature = (len(text_stripped) < 30 and 
                          any(s in text_stripped for s in ['部', '司', '局', '委', '办', '公司', '集团', '所', '院', '中心', '组', '室']))
            is_date = re.match(r'^\d{4}年\d{1,2}月\d{1,2}日$', text_stripped) or \
                     re.match(r'^[一二三四五六七八九十]{4}年[一二三四五六七八九十]{1,2}月[一二三四五六七八九十]{1,2}日$', text_stripped)
            
            if is_signature or is_date:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                para.paragraph_format.first_line_indent = Cm(0)
                chinese_font = '仿宋_GB2312'
                is_title = True
                para.add_run(original_text)
        
        # 13. 正文
        if not is_title:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.first_line_indent = Cm(1.27)
            chinese_font = '仿宋_GB2312'
            
            # 检查是否需要在结尾加句号
            is_date = re.match(r'^\d{4}年\d{1,2}月\d{1,2}日$', text_stripped) or \
                     re.match(r'^[一二三四五六七八九十]{4}年[一二三四五六七八九十]{1,2}月[一二三四五六七八九十]{1,2}日$', text_stripped)
            # 括号结尾不添加句号（如联系人信息）
            if text_stripped.endswith(')') or text_stripped.endswith('）'):
                pass
            elif text_stripped and not text_stripped.endswith(('。', '！', '？', '：', '；', '，', '日')) and len(text_stripped) > 3 and not is_date:
                original_text += '。'
            
            para.add_run(original_text)
        
        # 设置字体（中英文分离处理）
        # 先对整个original_text进行引号修正
        original_text = fix_quotes(original_text)
        for run in para.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = False  # 先清除所有加粗
            
        # 清空后重新添加（已经修正了引号）
        for run in para.runs:
            run.text = ''
        
        # 拆分中英文并分别设置字体
        parts = split_text(original_text)
        if len(parts) == 1:
            # 只有一种类型的文字
            text_part, is_ascii = parts[0]
            if is_ascii:
                run = para.add_run(text_part)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = False
            else:
                run = para.add_run(text_part)
                run.font.name = chinese_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = False
        else:
            # 需要拆分（中英文混合）
            for text_part, is_ascii in parts:
                new_run = para.add_run(text_part)
                new_run.font.size = Pt(16)
                new_run.font.color.rgb = RGBColor(0, 0, 0)
                new_run.bold = False
                if is_ascii:
                    new_run.font.name = 'Times New Roman'
                    new_run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                else:
                    new_run.font.name = chinese_font
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
    
    # 格式化表格
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    b = OxmlElement(f'w:{bn}')
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), '4')
                    b.set(qn('w:color'), '000000')
                    tcBorders.append(b)
                tcPr.append(tcBorders)
                
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = Pt(28)
                    for r in p.runs:
                        r.font.name = '仿宋_GB2312'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        r.font.size = Pt(16)
    
    doc.save(output_path)
    print(f"✅ 文档格式化完成：{output_path}")
    print("✅ 格式规范：")
    print("   • 一级标题：黑体三号，不加粗，缩进1.27cm")
    print("   • 三级标题：仿宋_GB2312三号，缩进1.27cm")
    print("   • 数字/英文：Times New Roman三号")
    print("   • 日期后面不加句号")
    print("   • 保留原文档空行结构，只格式化不改变段落分隔")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("用法：python format_docx.py 输入文件.docx 输出文件.docx 文档类型")
        sys.exit(1)
    
    try:
        format_official_document(sys.argv[1], sys.argv[2], sys.argv[3])
    except Exception as e:
        print(f"❌ 格式化失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
